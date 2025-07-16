"""
Unified document processor that creates a single vector store for all documents
with source metadata for proper attribution in responses.
"""
import os
import hashlib
from pathlib import Path
from typing import List, Tuple, Dict, Optional
import json
from datetime import datetime

from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    CSVLoader
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS

from src.config import Config
from src.local_llm import OptimizedLLM


class UnifiedDocumentProcessor:
    """Processes multiple documents into a single unified vector store"""
    
    UNIFIED_STORE_ID = "unified_store"
    
    def __init__(self):
        self.config = Config()
        self.config.create_directories()
        
        # Lazy initialization - don't require LLM just to load vector stores
        self.llm_system = None
        self.embeddings = None
        self.text_splitter = None
        
        # File type mappings (reuse from original document processor)
        self.file_loaders = {
            'pdf': PyPDFLoader,
            'txt': TextLoader,
            'text': TextLoader,
            'csv': CSVLoader,
            'md': TextLoader,  # Simplified for now
            'markdown': TextLoader,
            'docx': self._load_docx,
            'xlsx': self._load_xlsx,
            'png': self._load_image,
            'jpg': self._load_image,
            'jpeg': self._load_image
        }
    
    def _ensure_llm_initialized(self):
        """Ensure LLM and embeddings are initialized (lazy loading)"""
        if self.llm_system is None:
            self.llm_system = OptimizedLLM(self.config)
            self.embeddings = self.llm_system.get_embeddings()
            
        if self.text_splitter is None:
            optimal = self.config.get_optimal_settings()
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=optimal["chunk_size"],
                chunk_overlap=self.config.CHUNK_OVERLAP,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
    
    def _get_embeddings_for_loading(self):
        """Get embeddings for loading vector stores without requiring LLM"""
        # Create embeddings directly without LLM for loading vector stores
        from langchain_huggingface import HuggingFaceEmbeddings
        from src.memory_safe_embeddings import MemorySafeEmbeddings
        import os
        
        os.environ["HF_HUB_OFFLINE"] = "1"
        os.environ["TRANSFORMERS_OFFLINE"] = "1"
        
        try:
            base_embeddings = HuggingFaceEmbeddings(
                model_name=self.config.EMBEDDING_MODEL,
                model_kwargs={
                    'device': 'cpu',
                    'trust_remote_code': False,
                    'local_files_only': True
                },
                encode_kwargs={'normalize_embeddings': True}
            )
        except Exception:
            base_embeddings = HuggingFaceEmbeddings(
                model_name=self.config.EMBEDDING_MODEL,
                model_kwargs={'device': 'cpu', 'trust_remote_code': False},
                encode_kwargs={'normalize_embeddings': True}
            )
        
        return MemorySafeEmbeddings(
            base_embeddings, 
            batch_size=getattr(self.config, 'EMBEDDING_BATCH_SIZE', 2)
        )
    
    def detect_file_type(self, filename: str) -> str:
        """Detect file type from filename extension"""
        ext = filename.lower().split('.')[-1]
        
        type_mapping = {
            'pdf': 'pdf',
            'txt': 'text',
            'csv': 'csv',
            'md': 'markdown',
            'markdown': 'markdown',
            'docx': 'docx',
            'xlsx': 'xlsx',
            'png': 'png',
            'jpg': 'jpg',
            'jpeg': 'jpg'
        }
        
        if ext not in type_mapping:
            raise ValueError(f"Unsupported file type: .{ext}")
        
        return type_mapping[ext]
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """Custom Word document loader"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            import subprocess
            subprocess.check_call(["pip", "install", "python-docx"])
            from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        
        # Extract all paragraphs and tables
        content_blocks = []
        
        # Process paragraphs
        current_block = []
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    if current_block:
                        content_blocks.append('\n'.join(current_block))
                        current_block = []
                    current_block.append(para.text)
                else:
                    current_block.append(para.text)
        
        # Add the last block
        if current_block:
            content_blocks.append('\n'.join(current_block))
        
        # Process tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(' | '.join(row_text))
            
            if table_text:
                content_blocks.append('\n'.join(table_text))
        
        # Create documents from content blocks
        documents = []
        for i, block in enumerate(content_blocks):
            if block.strip():
                doc = Document(
                    page_content=block,
                    metadata={
                        "source": file_path,
                        "block_number": i,
                        "type": "docx"
                    }
                )
                documents.append(doc)
        
        return documents if documents else [Document(page_content="Empty document", metadata={"source": file_path})]
    
    def _load_xlsx(self, file_path: str) -> List[Document]:
        """Custom Excel loader that processes all sheets"""
        try:
            import openpyxl
        except ImportError:
            import subprocess
            subprocess.check_call(["pip", "install", "openpyxl"])
            import openpyxl
        
        documents = []
        
        # Load the workbook
        wb = openpyxl.load_workbook(file_path, data_only=True)
        
        # Process each sheet
        for sheet_idx, sheet_name in enumerate(wb.sheetnames):
            sheet = wb[sheet_name]
            
            # Skip empty sheets
            if sheet.max_row == 0 or sheet.max_column == 0:
                continue
            
            # Extract content from the sheet
            sheet_content = []
            sheet_content.append(f"Sheet: {sheet_name}")
            sheet_content.append("=" * 50)
            
            # Process rows
            for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
                # Filter out completely empty rows
                row_values = [str(cell) if cell is not None else '' for cell in row]
                if any(val.strip() for val in row_values):
                    # Format as: Row N: value1 | value2 | value3
                    row_text = f"Row {row_idx}: " + " | ".join(row_values)
                    sheet_content.append(row_text)
            
            # Create a document for this sheet
            if len(sheet_content) > 2:  # More than just header
                doc = Document(
                    page_content='\n'.join(sheet_content),
                    metadata={
                        "source": file_path,
                        "sheet_name": sheet_name,
                        "sheet_index": sheet_idx,
                        "type": "xlsx",
                        "total_rows": sheet.max_row,
                        "total_columns": sheet.max_column
                    }
                )
                documents.append(doc)
        
        return documents if documents else [Document(page_content="Empty spreadsheet", metadata={"source": file_path})]
    
    def _load_image(self, file_path: str) -> List[Document]:
        """Simplified image loader"""
        # For now, just create a placeholder
        doc = Document(
            page_content=f"Image file: {Path(file_path).name}",
            metadata={
                "source": file_path,
                "type": "image"
            }
        )
        return [doc]
    
    def process_documents(self, file_paths: List[str], progress_callback: callable = None) -> Dict[str, any]:
        """Process multiple documents into a unified vector store"""
        # Ensure LLM is initialized for processing
        self._ensure_llm_initialized()
        
        start_time = datetime.now()
        all_chunks = []
        document_metadata = []
        
        print(f"Processing {len(file_paths)} documents into unified vector store...")
        
        for idx, file_path in enumerate(file_paths):
            filename = Path(file_path).name
            print(f"\nProcessing {idx + 1}/{len(file_paths)}: {filename}")
            
            try:
                # Detect file type
                file_type = self.detect_file_type(filename)
                
                # Load documents
                if file_type in self.file_loaders:
                    loader_class = self.file_loaders[file_type]
                    
                    if callable(loader_class) and loader_class.__name__.startswith('_'):
                        # Custom loader method
                        documents = loader_class(file_path)
                    else:
                        # Standard LangChain loader
                        loader = loader_class(file_path)
                        documents = loader.load()
                else:
                    print(f"Skipping unsupported file type: {file_type}")
                    continue
                
                # Split into chunks
                chunks = self.text_splitter.split_documents(documents)
                
                # Add source metadata to each chunk
                for i, chunk in enumerate(chunks):
                    # Enhance metadata with source information
                    chunk.metadata.update({
                        'source_file': filename,
                        'source_path': file_path,
                        'file_type': file_type,
                        'chunk_index': i,
                        'total_chunks_in_file': len(chunks),
                        'page': chunk.metadata.get('page', chunk.metadata.get('row', 0))
                    })
                    
                    # Add source prefix to content for better context
                    chunk.page_content = f"[Source: {filename}]\n{chunk.page_content}"
                
                all_chunks.extend(chunks)
                
                # Track document metadata
                document_metadata.append({
                    'filename': filename,
                    'file_type': file_type,
                    'pages': len(documents),
                    'chunks': len(chunks),
                    'file_path': file_path
                })
                
                print(f"  ✓ Loaded {len(chunks)} chunks from {len(documents)} pages")
                
                if progress_callback:
                    progress_callback(idx + 1, len(file_paths), filename)
                    
            except Exception as e:
                print(f"  ✗ Error processing {filename}: {e}")
                continue
        
        if not all_chunks:
            raise ValueError("No documents were successfully processed")
        
        print(f"\nCreating unified embeddings for {len(all_chunks)} total chunks...")
        
        # Process embeddings in batches
        batch_size = self.config.BATCH_SIZE
        all_embeddings = []
        
        for i in range(0, len(all_chunks), batch_size):
            batch = all_chunks[i:i + batch_size]
            batch_texts = [chunk.page_content for chunk in batch]
            
            # Generate embeddings
            batch_embeddings = self.embeddings.embed_documents(batch_texts)
            all_embeddings.extend(batch_embeddings)
            
            # Progress indicator
            progress = min(100, (i + batch_size) / len(all_chunks) * 100)
            print(f"Embedding progress: {progress:.1f}%")
        
        # Create unified FAISS index
        print("Building unified vector index...")
        vector_store = FAISS.from_embeddings(
            text_embeddings=list(zip([c.page_content for c in all_chunks], all_embeddings)),
            embedding=self.embeddings,
            metadatas=[c.metadata for c in all_chunks]
        )
        
        # Save unified vector store
        vector_store_path = self.config.VECTOR_STORE_DIR / f"{self.UNIFIED_STORE_ID}.faiss"
        vector_store.save_local(str(vector_store_path))
        
        # Save unified metadata
        processing_time = (datetime.now() - start_time).total_seconds()
        metadata = {
            'store_id': self.UNIFIED_STORE_ID,
            'documents': document_metadata,
            'total_documents': len(document_metadata),
            'total_chunks': len(all_chunks),
            'creation_date': datetime.now().isoformat(),
            'processing_time': processing_time,
            'model_used': self.config.LOCAL_LLM_MODEL
        }
        
        metadata_path = self.config.VECTOR_STORE_DIR / f"{self.UNIFIED_STORE_ID}.metadata"
        with open(metadata_path, 'w') as f:
            json.dump(metadata, f, indent=2)
        
        print(f"\n✅ Unified vector store created successfully!")
        print(f"   - Documents: {len(document_metadata)}")
        print(f"   - Total chunks: {len(all_chunks)}")
        print(f"   - Processing time: {processing_time:.1f}s")
        
        return metadata
    
    def load_unified_store(self) -> FAISS:
        """Load the unified vector store"""
        vector_store_path = self.config.VECTOR_STORE_DIR / f"{self.UNIFIED_STORE_ID}.faiss"
        if not vector_store_path.exists():
            raise ValueError("Unified vector store not found. Please process documents first.")
        
        # Use lightweight embeddings for loading (no LLM required)
        embeddings = self._get_embeddings_for_loading()
        
        # Load vector store without signal-based timeout (doesn't work in threads)
        vector_store = FAISS.load_local(
            str(vector_store_path),
            embeddings,
            allow_dangerous_deserialization=True
        )
        return vector_store
    
    def get_unified_metadata(self) -> Optional[Dict]:
        """Get metadata for the unified store"""
        metadata_path = self.config.VECTOR_STORE_DIR / f"{self.UNIFIED_STORE_ID}.metadata"
        if not metadata_path.exists():
            return None
        
        with open(metadata_path, 'r') as f:
            return json.load(f)