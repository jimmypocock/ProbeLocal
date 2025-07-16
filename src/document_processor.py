import os
import hashlib
from pathlib import Path
from typing import List, Tuple, Dict
import pickle
import json
import time
from datetime import datetime
import csv
import pandas as pd

from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    CSVLoader
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
import numpy as np

from src.config import Config
from src.local_llm import OptimizedLLM
from src.incremental_processor import IncrementalProcessor
from src.security import validate_vector_store_path

class DocumentProcessor:
    def __init__(self):
        self.config = Config()
        self.config.create_directories()

        # Lazy initialization - don't require LLM just to load vector stores
        self.llm_system = None
        self.embeddings = None
        self.text_splitter = None
        
        # Initialize incremental processor for large files
        self.incremental_processor = IncrementalProcessor()
        
        # File size threshold for incremental processing (10MB)
        self.LARGE_FILE_THRESHOLD = 10 * 1024 * 1024
        
        # File type mappings
        self.file_loaders = {
            'pdf': PyPDFLoader,
            'txt': TextLoader,
            'text': TextLoader,
            'csv': self._load_csv,  # Custom CSV loader for better formatting
            'md': self._load_markdown,  # Custom Markdown loader
            'markdown': self._load_markdown,
            'docx': self._load_docx,  # Custom Word loader
            'pptx': self._load_pptx,  # TODO: Custom PowerPoint loader
            'xlsx': self._load_xlsx,  # Custom Excel loader
            'xls': self._load_xlsx,
            'png': self._load_image,  # Custom image loader with OCR
            'jpg': self._load_image,  # Custom image loader with OCR
            'jpeg': self._load_image
        }

    def _ensure_llm_initialized(self):
        """Ensure LLM and embeddings are initialized (lazy loading)"""
        if self.llm_system is None:
            self.llm_system = OptimizedLLM(self.config)
            self.embeddings = self.llm_system.get_embeddings()
            
        if self.text_splitter is None:
            optimal = self.config.get_optimal_settings()
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=optimal["chunk_size"],
                chunk_overlap=self.config.CHUNK_OVERLAP,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
    
    def _get_embeddings_for_loading(self):
        """Get embeddings for loading vector stores without requiring LLM"""
        # Create embeddings directly without LLM for loading vector stores
        from langchain_huggingface import HuggingFaceEmbeddings
        from src.memory_safe_embeddings import MemorySafeEmbeddings
        import os
        
        os.environ["HF_HUB_OFFLINE"] = "1"
        os.environ["TRANSFORMERS_OFFLINE"] = "1"
        
        try:
            base_embeddings = HuggingFaceEmbeddings(
                model_name=self.config.EMBEDDING_MODEL,
                model_kwargs={
                    'device': 'cpu',
                    'trust_remote_code': False,
                    'local_files_only': True
                },
                encode_kwargs={'normalize_embeddings': True}
            )
        except Exception:
            base_embeddings = HuggingFaceEmbeddings(
                model_name=self.config.EMBEDDING_MODEL,
                model_kwargs={'device': 'cpu', 'trust_remote_code': False},
                encode_kwargs={'normalize_embeddings': True}
            )
        
        return MemorySafeEmbeddings(
            base_embeddings, 
            batch_size=getattr(self.config, 'EMBEDDING_BATCH_SIZE', 2)
        )

    def detect_file_type(self, filename: str) -> str:
        """Detect file type from filename extension"""
        ext = filename.lower().split('.')[-1]
        
        # Map extensions to file types
        type_mapping = {
            'pdf': 'pdf',
            'txt': 'text',
            'csv': 'csv',
            'md': 'markdown',
            'markdown': 'markdown',
            'docx': 'docx',
            'pptx': 'pptx',
            'xlsx': 'xlsx',
            'xls': 'xlsx',
            'png': 'png',
            'jpg': 'jpg',
            'jpeg': 'jpg'
        }
        
        if ext not in type_mapping:
            raise ValueError(f"Unsupported file type: .{ext}")
        
        return type_mapping[ext]

    def _load_csv(self, file_path: str) -> List[Document]:
        """Custom CSV loader that preserves structure better"""
        documents = []
        
        try:
            # Read CSV file
            df = pd.read_csv(file_path)
            
            # Convert each row to a document
            for idx, row in df.iterrows():
                # Create a formatted text representation of the row
                content_parts = []
                for col, value in row.items():
                    if pd.notna(value):  # Skip NaN values
                        content_parts.append(f"{col}: {value}")
                
                content = "\n".join(content_parts)
                
                # Create document with metadata
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": file_path,
                        "row": idx,
                        "type": "csv"
                    }
                )
                documents.append(doc)
                
        except Exception as e:
            print(f"Error loading CSV: {e}")
            # Fallback to simple CSVLoader
            loader = CSVLoader(file_path)
            documents = loader.load()
        
        return documents
    
    def _load_markdown(self, file_path: str) -> List[Document]:
        """Custom Markdown loader"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Split by major sections (headers)
        sections = []
        current_section = []
        
        for line in content.split('\n'):
            if line.startswith('#') and current_section:
                # New section, save the previous one
                sections.append('\n'.join(current_section))
                current_section = [line]
            else:
                current_section.append(line)
        
        # Don't forget the last section
        if current_section:
            sections.append('\n'.join(current_section))
        
        # Create documents from sections
        documents = []
        for i, section in enumerate(sections):
            if section.strip():  # Skip empty sections
                doc = Document(
                    page_content=section,
                    metadata={
                        "source": file_path,
                        "section": i,
                        "type": "markdown"
                    }
                )
                documents.append(doc)
        
        # If no sections found, treat entire file as one document
        if not documents:
            doc = Document(
                page_content=content,
                metadata={
                    "source": file_path,
                    "type": "markdown"
                }
            )
            documents.append(doc)
        
        return documents
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """Custom Word document loader"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            print("Installing python-docx...")
            import subprocess
            subprocess.check_call(["pip", "install", "python-docx"])
            from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        
        # Extract all paragraphs and tables
        content_blocks = []
        
        # Process paragraphs
        current_block = []
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    if current_block:
                        content_blocks.append('\n'.join(current_block))
                        current_block = []
                    current_block.append(para.text)
                else:
                    current_block.append(para.text)
        
        # Add the last block
        if current_block:
            content_blocks.append('\n'.join(current_block))
        
        # Process tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                table_text.append(' | '.join(row_text))
            if table_text:
                content_blocks.append('\n'.join(table_text))
        
        # Create documents from content blocks
        documents = []
        for i, block in enumerate(content_blocks):
            if block.strip():
                doc = Document(
                    page_content=block,
                    metadata={
                        "source": file_path,
                        "block": i,
                        "type": "docx"
                    }
                )
                documents.append(doc)
        
        # If no content found, create a single empty document
        if not documents:
            doc = Document(
                page_content="Empty document",
                metadata={
                    "source": file_path,
                    "type": "docx"
                }
            )
            documents.append(doc)
        
        return documents
    
    def _load_pptx(self, file_path: str) -> List[Document]:
        """Placeholder for PowerPoint loader"""
        raise NotImplementedError("PowerPoint support coming soon")
    
    def _load_xlsx(self, file_path: str) -> List[Document]:
        """Custom Excel loader that processes all sheets"""
        try:
            import openpyxl
        except ImportError:
            print("Installing openpyxl...")
            import subprocess
            subprocess.check_call(["pip", "install", "openpyxl"])
            import openpyxl
        
        documents = []
        
        # Load the workbook
        wb = openpyxl.load_workbook(file_path, data_only=True)
        
        # Process each sheet
        for sheet_idx, sheet_name in enumerate(wb.sheetnames):
            sheet = wb[sheet_name]
            
            # Skip empty sheets
            if sheet.max_row == 0 or sheet.max_column == 0:
                continue
            
            # Extract content from the sheet
            sheet_content = []
            sheet_content.append(f"Sheet: {sheet_name}")
            sheet_content.append("=" * 50)
            
            # Process rows
            for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
                # Filter out completely empty rows
                if all(cell is None for cell in row):
                    continue
                
                # Format row content
                row_text = []
                for col_idx, cell_value in enumerate(row):
                    if cell_value is not None:
                        # Convert to string and clean up
                        cell_str = str(cell_value).strip()
                        if cell_str:
                            # If it's the first row and looks like headers, format differently
                            if row_idx == 1 and col_idx == 0:
                                row_text.append(cell_str)
                            else:
                                row_text.append(cell_str)
                
                if row_text:
                    # Join cells with appropriate separator
                    if len(row_text) > 1:
                        sheet_content.append(" | ".join(row_text))
                    else:
                        sheet_content.append(row_text[0])
            
            # Create a document for this sheet
            if len(sheet_content) > 2:  # More than just header
                content = "\n".join(sheet_content)
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": file_path,
                        "sheet_name": sheet_name,
                        "sheet_index": sheet_idx,
                        "type": "xlsx"
                    }
                )
                documents.append(doc)
        
        # If no content found, create a single document
        if not documents:
            doc = Document(
                page_content="Empty Excel file",
                metadata={
                    "source": file_path,
                    "type": "xlsx"
                }
            )
            documents.append(doc)
        
        return documents
    
    def _load_image(self, file_path: str) -> List[Document]:
        """Custom image loader with OCR text extraction"""
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            print("Installing OCR dependencies...")
            import subprocess
            subprocess.check_call(["pip", "install", "pytesseract", "Pillow"])
            import pytesseract
            from PIL import Image
        
        documents = []
        
        try:
            # Load and process the image
            with Image.open(file_path) as image:
                # Extract text using OCR
                extracted_text = pytesseract.image_to_string(image)
            
            # Clean up the extracted text
            lines = []
            for line in extracted_text.split('\n'):
                line = line.strip()
                if line:  # Skip empty lines
                    lines.append(line)
            
            # Create content
            if lines:
                content = '\n'.join(lines)
                
                # Add image description header
                file_ext = file_path.lower().split('.')[-1]
                content = f"Image Content (OCR Extracted Text):\n{'-'*40}\n{content}"
            else:
                # No text found - this is likely a visual content image
                file_ext = file_path.lower().split('.')[-1]
                content = f"Visual Image Content ({file_ext.upper()}):\n{'-'*40}\n"
                content += "This appears to be a visual image with no extractable text content. "
                content += "The image contains visual elements that would require computer vision analysis "
                content += "to describe. This could include objects, people, animals, scenery, or other "
                content += "visual content that cannot be extracted through OCR text recognition."
            
            # Create document
            doc = Document(
                page_content=content,
                metadata={
                    "source": file_path,
                    "type": file_ext,
                    "extraction_method": "OCR",
                    "image_format": file_ext.upper()
                }
            )
            documents.append(doc)
            
        except Exception as e:
            print(f"Error processing image: {e}")
            # Create a basic document even if OCR fails
            file_ext = file_path.lower().split('.')[-1]
            doc = Document(
                page_content=f"Image file ({file_ext.upper()}) - OCR processing failed",
                metadata={
                    "source": file_path,
                    "type": file_ext,
                    "extraction_method": "failed",
                    "error": str(e)
                }
            )
            documents.append(doc)
        
        return documents

    def generate_document_id(self, file_path: str) -> str:
        """Generate unique ID for document"""
        with open(file_path, 'rb') as f:
            file_hash = hashlib.sha256(f.read()).hexdigest()
        return file_hash[:16]

    def process_file(self, file_path: str, filename: str, chunk_size: int = None, progress_callback: callable = None) -> Tuple[str, int, int, float]:
        """Process any supported file type with optional dynamic chunk size"""
        # Ensure LLM is initialized for processing
        self._ensure_llm_initialized()
        
        start_time = time.time()
        
        # Check file size for incremental processing
        file_size = os.path.getsize(file_path)
        
        # Use incremental processing for large files
        if file_size > self.LARGE_FILE_THRESHOLD:
            print(f"Large file detected ({file_size / 1024 / 1024:.1f} MB), using incremental processing...")
            return self.incremental_processor.process_file_incremental(
                file_path=file_path,
                file_name=filename,
                chunk_size=chunk_size,
                progress_callback=progress_callback
            )
        
        # Regular processing for smaller files
        # Detect file type
        file_type = self.detect_file_type(filename)
        
        print(f"Loading {file_type.upper()} file: {filename}")
        
        # Load documents based on file type
        if file_type in self.file_loaders:
            loader_class = self.file_loaders[file_type]
            
            if callable(loader_class) and loader_class.__name__.startswith('_'):
                # Custom loader method
                documents = loader_class(file_path)
            else:
                # Standard LangChain loader
                loader = loader_class(file_path)
                documents = loader.load()
        else:
            raise ValueError(f"No loader available for file type: {file_type}")

        doc_id = self.generate_document_id(file_path)

        # Use dynamic chunk size if provided, otherwise use default
        if chunk_size is not None:
            print(f"Splitting into chunks (custom size: {chunk_size})...")
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=self.config.CHUNK_OVERLAP,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
            chunks = text_splitter.split_documents(documents)
        else:
            print(f"Splitting into chunks (default size: {self.config.CHUNK_SIZE})...")
            chunks = self.text_splitter.split_documents(documents)

        # Handle empty files - create a minimal document if no chunks exist
        if not chunks:
            print("Warning: Empty file detected, creating minimal document")
            empty_doc = Document(
                page_content=f"Empty {file_type.upper()} file: {filename}",
                metadata={
                    'document_id': doc_id,
                    'filename': filename,
                    'file_type': file_type,
                    'chunk_index': 0,
                    'total_chunks': 1,
                    'page': 0,
                    'is_empty': True
                }
            )
            chunks = [empty_doc]

        # Add metadata
        for i, chunk in enumerate(chunks):
            chunk.metadata.update({
                'document_id': doc_id,
                'filename': filename,
                'file_type': file_type,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'page': chunk.metadata.get('page', chunk.metadata.get('row', 0))
            })

        print(f"Creating embeddings for {len(chunks)} chunks...")

        # Process in batches to manage memory
        batch_size = self.config.BATCH_SIZE
        all_embeddings = []

        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            batch_texts = [chunk.page_content for chunk in batch]

            # Generate embeddings
            batch_embeddings = self.embeddings.embed_documents(batch_texts)
            all_embeddings.extend(batch_embeddings)

            # Progress indicator
            progress = min(100, (i + batch_size) / len(chunks) * 100)
            print(f"Progress: {progress:.1f}%")

        # Create FAISS index
        print("Building vector index...")
        vector_store = FAISS.from_embeddings(
            text_embeddings=list(zip([c.page_content for c in chunks], all_embeddings)),
            embedding=self.embeddings,
            metadatas=[c.metadata for c in chunks]
        )

        # Save vector store
        vector_store_path = self.config.VECTOR_STORE_DIR / f"{doc_id}.faiss"
        vector_store.save_local(str(vector_store_path))

        # Save metadata
        processing_time = time.time() - start_time
        metadata = {
            'document_id': doc_id,
            'filename': filename,
            'file_type': file_type,
            'pages': len(documents),
            'chunks': len(chunks),
            'upload_date': datetime.now(),
            'processing_time': processing_time,
            'file_size_mb': os.path.getsize(file_path) / (1024 * 1024),
            'model_used': self.config.LOCAL_LLM_MODEL
        }

        metadata_path = self.config.VECTOR_STORE_DIR / f"{doc_id}.metadata"
        # Convert datetime to string for JSON serialization
        metadata_json = metadata.copy()
        metadata_json['upload_date'] = metadata_json['upload_date'].isoformat()
        with open(metadata_path, 'w') as f:
            json.dump(metadata_json, f)

        print(f"Processing complete in {processing_time:.1f} seconds")
        
        # Auto-cleanup old documents if needed
        self._cleanup_old_documents()
        
        return doc_id, len(documents), len(chunks), processing_time

    def process_pdf(self, file_path: str, filename: str) -> Tuple[str, int, int, float]:
        """Legacy method for backward compatibility - redirects to process_file"""
        return self.process_file(file_path, filename)

    def load_vector_store(self, document_id: str) -> FAISS:
        """Load existing vector store"""
        vector_store_path = self.config.VECTOR_STORE_DIR / f"{document_id}.faiss"
        if not vector_store_path.exists():
            raise ValueError(f"Document {document_id} not found")

        # Validate the vector store path for security
        if not validate_vector_store_path(vector_store_path, self.config.VECTOR_STORE_DIR):
            raise ValueError(f"Invalid vector store path for document ID: {document_id}")

        # Use lightweight embeddings for loading (no LLM required)
        embeddings = self._get_embeddings_for_loading()

        # Load with safer approach - still uses pickle internally but with validation
        # In production, consider migrating to a safer serialization format
        return FAISS.load_local(
            str(vector_store_path),
            embeddings,
            allow_dangerous_deserialization=True  # Required by FAISS, but path is validated
        )
    
    def _cleanup_old_documents(self):
        """Remove old documents to manage storage"""
        try:
            max_docs = int(os.getenv('MAX_DOCUMENTS', '20'))
            cleanup_days = int(os.getenv('CLEANUP_DAYS', '7'))
            
            # Get all metadata files with their modification times
            metadata_files = []
            for f in self.config.VECTOR_STORE_DIR.glob("*.metadata"):
                metadata_files.append((f, f.stat().st_mtime))
            
            # Sort by modification time (oldest first)
            metadata_files.sort(key=lambda x: x[1])
            
            # Remove oldest documents if we exceed max count
            if len(metadata_files) > max_docs:
                files_to_remove = len(metadata_files) - max_docs
                for i in range(files_to_remove):
                    metadata_file = metadata_files[i][0]
                    doc_id = metadata_file.stem
                    
                    # Remove vector store
                    vector_store_path = self.config.VECTOR_STORE_DIR / f"{doc_id}.faiss"
                    if vector_store_path.exists():
                        import shutil
                        shutil.rmtree(vector_store_path)
                    
                    # Remove metadata
                    metadata_file.unlink()
                    
                    print(f"Auto-removed old document: {doc_id}")
            
            # Also remove documents older than cleanup_days
            current_time = time.time()
            for metadata_file, mtime in metadata_files:
                age_days = (current_time - mtime) / (24 * 3600)
                if age_days > cleanup_days:
                    doc_id = metadata_file.stem
                    
                    # Remove vector store
                    vector_store_path = self.config.VECTOR_STORE_DIR / f"{doc_id}.faiss"
                    if vector_store_path.exists():
                        import shutil
                        shutil.rmtree(vector_store_path)
                    
                    # Remove metadata
                    metadata_file.unlink()
                    
                    print(f"Auto-removed expired document: {doc_id} (age: {age_days:.1f} days)")
                    
        except Exception as e:
            print(f"Cleanup error: {e}")
            # Don't fail the main process if cleanup fails
            pass